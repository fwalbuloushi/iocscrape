import gzip
import os
import re
import urllib.error
import urllib.request
import zlib
from typing import Dict, Tuple

import trafilatura

from .constants import APP_NAME, APP_VERSION, PROJECT_URL


def _decompress(data: bytes, content_encoding: str) -> bytes:
    enc = (content_encoding or "").lower()
    if "gzip" in enc:
        return gzip.decompress(data)
    if "deflate" in enc:
        try:
            return zlib.decompress(data)
        except zlib.error:
            return zlib.decompress(data, -zlib.MAX_WBITS)
    return data


def fetch_url_bytes(
    url: str, timeout: int, max_bytes: int, redirect_limit: int
) -> Tuple[bytes, str, Dict[str, str]]:
    current_url = url
    accept_encoding = "gzip, deflate, identity"

    for _ in range(redirect_limit + 1):
        req = urllib.request.Request(
            current_url,
            headers={
                "User-Agent": f"{APP_NAME}/{APP_VERSION} (+{PROJECT_URL})",
                "Accept-Encoding": accept_encoding,
            },
            method="GET",
        )

        try:
            with urllib.request.urlopen(req, timeout=timeout) as resp:
                status = getattr(resp, "status", None)
                final_url = resp.geturl()
                headers = {k.lower(): v for k, v in resp.headers.items()}

                if status in (301, 302, 303, 307, 308):
                    loc = resp.headers.get("Location")
                    if not loc:
                        break
                    current_url = urllib.request.urljoin(final_url, loc)
                    continue

                data = resp.read(max_bytes + 1)
                if len(data) > max_bytes:
                    raise ValueError(f"Response exceeded max size ({max_bytes} bytes).")

                ce = headers.get("content-encoding", "")
                if "br" in (ce or "").lower() and accept_encoding != "identity":
                    accept_encoding = "identity"
                    continue

                data = _decompress(data, ce)
                return data, final_url, headers

        except urllib.error.HTTPError as e:
            raise RuntimeError(f"HTTP error fetching URL: {e.code} {e.reason}") from e
        except urllib.error.URLError as e:
            raise RuntimeError(f"Network error fetching URL: {e.reason}") from e

    raise RuntimeError(f"Too many redirects (limit={redirect_limit}).")


def decode_bytes(data: bytes, headers: Dict[str, str]) -> str:
    ct = headers.get("content-type", "")
    m = re.search(r"charset=([^\s;]+)", ct, re.IGNORECASE)
    if m:
        cs = m.group(1).strip()
        try:
            return data.decode(cs, errors="replace")
        except LookupError:
            pass
    return data.decode("utf-8", errors="replace")


def read_file_text(path: str) -> str:
    if not os.path.isfile(path):
        raise FileNotFoundError(f"File not found: {path}")

    ext = os.path.splitext(path)[1].lower().strip(".")

    if ext in ("txt", "html", "htm"):
        with open(path, "rb") as f:
            return f.read().decode("utf-8", errors="replace")

    if ext == "pdf":
        try:
            from pdfminer.high_level import extract_text as pdf_extract_text
        except Exception as e:
            raise RuntimeError("pdfminer.six is required for PDF parsing.") from e
        return pdf_extract_text(path) or ""

    if ext == "docx":
        try:
            import docx
        except Exception as e:
            raise RuntimeError("python-docx is required for DOCX parsing.") from e
        d = docx.Document(path)
        parts = []
        for p in d.paragraphs:
            if p.text:
                parts.append(p.text)
        for table in d.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    t = (cell.text or "").strip()
                    if t:
                        row_text.append(t)
                if row_text:
                    parts.append(" | ".join(row_text))
        return "\n".join(parts)

    if ext == "xlsx":
        try:
            import openpyxl
        except Exception as e:
            raise RuntimeError("openpyxl is required for XLSX parsing.") from e
        wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
        parts = []
        for sheet in wb.worksheets:
            parts.append(f"[SHEET] {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(v).strip() for v in row if v is not None and str(v).strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    raise RuntimeError(f"Unsupported file extension: .{ext} (supported: txt, html, pdf, docx, xlsx)")


def extract_article_text_from_html(html: str) -> str:
    text = trafilatura.extract(
        html,
        output_format="txt",
        include_links=False,
        include_images=False,
        include_tables=True,
        favor_precision=True,
        deduplicate=True,
    )
    return (text or "").strip()


def looks_like_html(s: str) -> bool:
    sl = s.lower()
    return "<html" in sl or "</" in sl
